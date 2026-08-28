import open3d as o3d
import numpy as np
from sklearn.neighbors import KDTree
import pylas
import random
from tqdm import tqdm
from pykdtree.kdtree import KDTree
import datetime

from scipy import spatial
import os
import sys
import laspy

from scipy.optimize import fsolve
from scipy.optimize import root

from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from scipy.optimize import curve_fit
from docx import Document
from docx.shared import Pt, RGBColor

starttime = None

fn = None

progress_value = 0

def read_file(file_name):
    global progress_value
    file = pylas.read(file_name)
    point_all = np.vstack((file.x,file.y,file.z)).transpose()
    pcd = o3d.geometry.PointCloud()
    pcd.points = o3d.utility.Vector3dVector(point_all)

    progress_value = 5
    return pcd

def filter_elevation(pcd):
    point = np.array(pcd.points)
    filtered_points = point[point[:, 2] >= 30]
    filtered_pcd = o3d.geometry.PointCloud()
    filtered_pcd.points = o3d.utility.Vector3dVector(filtered_points)
    return filtered_pcd

def simple(pcd, ratio=1):
    #points = np.asarray(pcd.points).astype(np.float32)
    #indices = np.random.choice(len(points), size=int(len(points) * ratio), replace=False)
    #down_pcd = pcd.select_by_index(indices)
    #return down_pcd
    
    # denoise
    dn_nb = 3
    dn_std = 5
    filter_pcd, _ = pcd.remove_statistical_outlier(nb_neighbors=dn_nb, std_ratio=dn_std)  ###6 3
    
    filter_points = np.asarray(filter_pcd.points)
    print('filter point cloud:', filter_points.shape)
    
    #np.savetxt("./results/" + fn + '/' + fn + '_denosie.xyz', filter_points, fmt = '%.6f')
    
    return filter_pcd
    
def rough_extraction(pcd_org):
    global progress_value
    
    #down sample to 1/4
    voxel_size = 1.8  #1/4
    
    down_sample_pcd = pcd_org.voxel_down_sample(voxel_size)
    
    down_sample_points = np.asarray(down_sample_pcd.points)
    
    print('downsample point cloud:', down_sample_points.shape)
    
    progress_value = 12
    print('step 3/8-1:',(datetime.datetime.now() - starttime))
    
    #np.savetxt("./results/" + fn + '/' + fn + '_sample.xyz', down_sample_points, fmt = '%.6f')
    
    pcd = down_sample_pcd
    
    #pcd = pcd_org

    points_org = np.asarray(pcd_org.points).astype(np.float32)
    
    points = np.asarray(pcd.points).astype(np.float32)
    kdtree = KDTree(points)
    
        
    #R = 3
    #knn = 30
    
    #R = 3.5
    #knn = 30
    
    #R = 5.0
    R = 5.3 
    knn = 30 
    
    #R = 5
    #knn = 100
    
    #R = 10
    #knn = 100
    
    print('计算邻居')
    
    num_org = len(points_org)
    num = len(points)
    
    ## batch 计算邻居
    progress_bar = tqdm(total=num_org, desc="邻居计算进度")
    step = 100000
    indices_all = []
    for i in range(0, num_org, step):
        batch_end = min(i + step, num_org)
        batch_points = points_org[i:batch_end]
        _, indices = kdtree.query(batch_points, k=knn, distance_upper_bound=R)
        indices_all.extend(indices)
        progress_bar.update(batch_end - i)
        
        progress_value = 12 + (35-12)*batch_end/num_org
        
    indices_all = np.array(indices_all).astype(np.int32)
    #print(indices_all.shape)
    progress_bar.close()
    
    
    ## all-one time 计算邻居
    '''
    _, indices_all = kdtree.query(points, k=knn, distance_upper_bound=R)
    indices_all = indices_all.astype(np.int32)
    '''
    progress_value = 35
    print('step 3/8-2:',(datetime.datetime.now() - starttime))
    
    
    ##修正溢出2
    print('修正邻域索引')
    
    #batch 修正
    '''
    step = 1000000
    num = len(points)
    knn_nums = []
    indices_all
    for i in tqdm(range(0,len(points), step)):
        indices_all_0_1 = np.where(indices_all[i: min(i+ step, num),] == num, 1, 0)
        knn_nums.extend(knn - np.sum(indices_all_0_1, axis = 1))
        indices_all[i: min(i+ step, num),] = indices_all[i: min(i+ step, num),] + indices_all_0_1 * (np.arange(-num, 0, 1).reshape(-1,1)[i: min(i+ step, num),])  ## self
    knn_nums = np.array(knn_nums).astype(np.int32)
    '''
    #all-one time 修正
    
    indices_all_0_1 = np.where(indices_all == num, 1, 0)
    knn_nums = knn - np.sum(indices_all_0_1, axis = 1)
    indices_all = indices_all + indices_all_0_1 * np.arange(0, num_org, 1).reshape(-1,1)
    #indices_all = indices_all + indices_all_0_1 * np.arange(-num, 0, 1).reshape(-1,1)  ## self
    
    progress_value = 36
    print('step 3/8-3:',(datetime.datetime.now() - starttime))
    
    ## all
    print('获取邻域点和协方差矩阵')
    linearity = np.zeros(num_org)
    planarity = np.zeros(num_org)
    scattering = np.zeros(num_org)
    valid_idx = []
    
    valid_idx = np.where(knn_nums>=4)[0]
    invalid_idx = np.where(knn_nums<4)[0]
    #print('有效点数和无效点数:', valid_idx.shape, invalid_idx.shape)
    linearity[invalid_idx] = 0
    planarity[invalid_idx] = 0
    scattering[invalid_idx] = 1
    
    # batch get knn
    ### num + num_org
    step = 100000
    extend_points = np.concatenate([points, points_org], axis = 0).astype(np.float32)
    cov_matrix = []
    for i in tqdm(range(0,len(valid_idx), step)):
        batch_end = min(i+ step, len(valid_idx))
        neighbors = extend_points[indices_all[valid_idx[i: batch_end]]]
        centered_points = neighbors - points_org[valid_idx[i: batch_end]].reshape(-1,1,3)
        cov_matrix_temp = np.matmul(centered_points.transpose(0,2,1), centered_points).astype(np.float32) / (knn_nums[valid_idx[i: batch_end]].reshape(-1,1,1) - 1)
        cov_matrix.append(cov_matrix_temp)
        
        progress_value = 36 + (44-36)*batch_end/len(valid_idx)
    cov_matrix = np.concatenate(cov_matrix, axis = 0)
    
    progress_value = 44
    print('step 3/8-4:',(datetime.datetime.now() - starttime))
    
    ############# svd
    print('协方差矩阵求解特征值')
    eigenvalues = np.linalg.eigvalsh(cov_matrix)
    eigenvalues = np.sort(eigenvalues, axis = -1)[:,::-1]
    eigenvalues = np.clip(eigenvalues, 1e-8, None)
    
    #print(eigenvalues[:10])
    
    
    lambda1, lambda2, lambda3 = eigenvalues[:,0],eigenvalues[:,1],eigenvalues[:,2]
    L_lambda = (lambda1 - lambda2) / lambda1
    P_lambda = (lambda2 - lambda3) / lambda1
    S_lambda = lambda3 / lambda1
    linearity[valid_idx] = L_lambda
    planarity[valid_idx] = P_lambda
    scattering[valid_idx] = S_lambda
    
    progress_value = 47
    print('step 3/8-5:',(datetime.datetime.now() - starttime))
    
    
    ############
    print('获取候选点')
    linearity_threshold = 0.8
    planarity_threshold = 0.6
    
    idx_c1 = np.where((linearity > linearity_threshold) & (planarity < planarity_threshold))[0]
    
    #colors = np.zeros([num_org, 3]) + [0,1,0]
    #colors[idx_c1] = [1, 0, 0]
    
    #np.savetxt('./results/' + fn + '/' + fn + '_linearity_planarity_' + str(R) + '_' + str(knn) + '.xyz', np.concatenate([points_org,colors,linearity.reshape(-1,1),planarity.reshape(-1,1)], axis = 1), fmt = '%.6f')
    
    red_points = points_org[idx_c1]
    red_pcd = o3d.geometry.PointCloud()
    red_pcd.points = o3d.utility.Vector3dVector(red_points)
    red_pcd.colors = o3d.utility.Vector3dVector(np.zeros([len(idx_c1), 3]) + [1,0,0])
    
    print('候选点数:',red_points.shape)
    
    return red_pcd

    
def DBSCAN(pcd, eps=3.5, min_points=10, print_progress=True):  # eps3.5
    # eps越小，越少点被加入核心
    if not isinstance(pcd, o3d.geometry.PointCloud):
        raise TypeError("输入必须是 Open3D 点云对象")
    pcd_result = pcd
    labels = np.array(pcd_result.cluster_dbscan(eps=eps, min_points=min_points, print_progress=print_progress))
    max_label = labels.max()
    n_clusters = max_label + 1
    n_noise = list(labels).count(-1)
    colors = np.zeros((len(pcd_result.points), 3))
    colors[labels < 0] = [0.5, 0.5, 0.5]
    pcds = []
    for i in range(n_clusters):
        rand_color = np.random.rand(3)
        colors[labels == i] = rand_color
        indices = np.where(labels == i)[0]
        cluster_pcd = o3d.geometry.PointCloud()
        cluster_pcd.points = o3d.utility.Vector3dVector(np.asarray(pcd_result.points)[indices])
        if pcd_result.has_normals():
            cluster_pcd.normals = o3d.utility.Vector3dVector(np.asarray(pcd_result.normals)[indices])
        cluster_colors = np.full((len(indices), 3), rand_color)
        cluster_pcd.colors = o3d.utility.Vector3dVector(cluster_colors)
        pcds.append(cluster_pcd)
    if n_noise > 0:
        noise_indices = np.where(labels == -1)[0]
        noise_pcd = o3d.geometry.PointCloud()
        noise_pcd.points = o3d.utility.Vector3dVector(np.asarray(pcd_result.points)[noise_indices])
        if pcd_result.has_normals():
            noise_pcd.normals = o3d.utility.Vector3dVector(np.asarray(pcd_result.normals)[noise_indices])
        noise_colors = np.full((len(noise_indices), 3), [0.5, 0.5, 0.5])
        noise_pcd.colors = o3d.utility.Vector3dVector(noise_colors)
    pcd_result.colors = o3d.utility.Vector3dVector(colors)
    print('\n初始线条数量:', len(pcds))
    return pcds

def maxpoint(pcd, radius=20):
    if not isinstance(pcd, o3d.geometry.PointCloud):
        raise TypeError("输入必须是open3d.geometry.PointCloud类型")
    points = np.asarray(pcd.points)
    if len(points) == 0:
        return [pcd]
    kdtree = o3d.geometry.KDTreeFlann(pcd)
    max_points_indices = []
    for i in range(len(points)):
        [k, idx, _] = kdtree.search_radius_vector_3d(points[i], radius)
        if k <= 1:
            continue
        z_values = points[idx, 2]
        current_z = points[i, 2]
        if current_z >= np.max(z_values):
            max_points_indices.append(i)
    if not max_points_indices:
        return [pcd]
    planes = []
    for max_idx in max_points_indices:
        max_point = points[max_idx]
        distances = np.linalg.norm(points - max_point, axis=1)
        furthest_idx = np.argmax(distances)
        furthest_point = points[furthest_idx]
        normal = furthest_point - max_point
        normal = normal / np.linalg.norm(normal)
        d = -np.dot(normal, max_point)
        planes.append((normal, d, max_point))
    segment_labels = np.zeros(len(points), dtype=int) - 1
    for i in range(len(points)):
        point = points[i]
        point_position = []
        for normal, d, _ in planes:
            distance = np.dot(normal, point) + d
            point_position.append(distance >= 0)
        label = 0
        for j, pos in enumerate(point_position):
            if pos:
                label |= (1 << j)
        segment_labels[i] = label
    unique_labels = np.unique(segment_labels)
    segment_pcds = []
    for label in unique_labels:
        segment_indices = np.where(segment_labels == label)[0]
        segment_points = points[segment_indices]
        segment_pcd = o3d.geometry.PointCloud()
        segment_pcd.points = o3d.utility.Vector3dVector(segment_points)
        random_color = [random.random(), random.random(), random.random()]
        segment_colors = np.tile(random_color, (len(segment_points), 1))
        segment_pcd.colors = o3d.utility.Vector3dVector(segment_colors)
        segment_pcds.append(segment_pcd)
    return segment_pcds

def maxpoints(pcds):
    global progress_value
    
    result_pcds = []
    for i, pcd in enumerate(pcds):
        divided = maxpoint(pcd)
        result_pcds.extend(divided)
        
        progress_value = 50 + (55-50)*i/len(pcds)
    return result_pcds

def is_line(pcds, simple_pcd, linearity_threshold=0.995, planarity_threshold=0.2):
    line_pcds = []
    for i, pcd in enumerate(pcds):
        points = np.asarray(pcd.points)
        if len(points) < 200:
            continue
        centroid = np.mean(points, axis=0)
        centered_points = points - centroid
        cov_matrix = np.cov(centered_points, rowvar=False)
        eigenvalues, eigenvectors = np.linalg.eigh(cov_matrix)
        eigenvalues = eigenvalues[::-1]
        eigenvectors = eigenvectors[:, ::-1]
        linearity = (eigenvalues[0] - eigenvalues[1]) / eigenvalues[0] if eigenvalues[0] > 0 else 0
        planarity = (eigenvalues[1] - eigenvalues[2]) / eigenvalues[0] if eigenvalues[0] > 0 else 0
        if linearity > linearity_threshold and planarity < planarity_threshold:
            line_pcds.append(pcd)
    print('筛选线条数量:',len(line_pcds))
    if len(line_pcds) == 0:
        return o3d.geometry.PointCloud()
        
    combined_pcd = o3d.geometry.PointCloud()
    for pcd in line_pcds:
        combined_pcd += pcd
    
    # fitting line
    towers_centers = fit_line(line_pcds, simple_pcd)
    
    points_all = np.asarray(combined_pcd.points)
    colors_all = np.asarray(combined_pcd.colors)
    #np.savetxt('./results/' + fn + '/' + fn + '_results_select.xyz', np.concatenate([points_all, colors_all*255],axis = 1), fmt = '%.6f')
    print('筛选线条的总点数:',points_all.shape)
    return combined_pcd, line_pcds, towers_centers
    
def fit_line(pcds, simple_pcd):
    starts = []
    ends = []
    for i, pcd in enumerate(pcds):
        points = np.asarray(pcd.points)
        start = points[np.argmin(points[:,0]),:]
        end = points[np.argmax(points[:,0]),:]
        
        starts.append(start)
        ends.append(end)
    starts = np.array(starts)
    ends = np.array(ends)
    
    #np.savetxt('./results/' + fn + '/' + fn + '_starts.xyz', starts)
    #np.savetxt('./results/' + fn + '/' + fn + '_ends.xyz', ends)
    
    
    points = np.asarray(simple_pcd.points)
    kdtree = spatial.KDTree(points[:,:2])
    #### add ta
    tower_init = np.loadtxt('./model/tower.xyz')
    tower_init[:,:2] = tower_init[:,:2] - np.mean(tower_init[:,:2], axis = 0).reshape(-1,2)
    tower_init[:,2:3] = tower_init[:,2:3] - np.min(tower_init[:,2:3]).reshape(-1,1)
    height_init = np.max(tower_init[:,2:3])
    
    pts = np.concatenate([starts,ends], axis = 0)
    pts_pcd = o3d.geometry.PointCloud()
    pts_pcd.points = o3d.utility.Vector3dVector(pts)
    
    labels = np.array(pts_pcd.cluster_dbscan(eps=35, min_points=4, print_progress=True))
    max_label = labels.max()
    #print('./max label:',max_label)
    n_clusters = max_label + 1
    n_noise = list(labels).count(-1)
    towers = []
    towers2 = []
    towers_orgs = []
    towers_centers = []
    for i in range(n_clusters):
        point_org = pts[np.where(labels == i)[0]]
        towers_orgs.append(point_org)
        
        towers_center = np.mean(point_org, axis = 0)
        idx = kdtree.query_ball_point(towers_center[:2],15)
        points_query = points[idx,]
        towers_centers.append(towers_center)
        
        pcd_temp = o3d.geometry.PointCloud()
        pcd_temp.points = o3d.utility.Vector3dVector(points_query)
        dn_nb = 10
        dn_std = 1
        cl, _ = pcd_temp.remove_statistical_outlier(nb_neighbors=dn_nb, std_ratio=dn_std)  ###6 3
        points_query_dn = np.asarray(cl.points)   
        
        towers2.append(points_query)
        
        temp_height = np.max(points_query_dn[:,2:3]) - np.min(points_query_dn[:,2:3])
        temp_min = np.min(points_query_dn[:,2:3])
        
        if(temp_height>80):
            temp_height = 80
            temp_min = np.max(points_query_dn[:,2:3]) - 80
            
        tower_temp = tower_init.copy() / height_init * temp_height * 1.05
        tower_temp = tower_temp+ [towers_center[0],towers_center[1], temp_min]
        
        towers.append(tower_temp)
        
    towers_centers = np.array(towers_centers)
    
    towers = np.concatenate(towers, axis = 0)
    colors = np.zeros_like(towers) + [[255, 255, 255]]
    #np.savetxt('./results/' + fn + '/' + fn + '_towers.xyz', towers, fmt = '%.6f')
    
    header = laspy.LasHeader(version="1.4", point_format=7)  # 点格式1对应XYZI，根据你的数据调整格式
    inFile = laspy.LasData(header)
    inFile.x = towers[:, 0]
    inFile.y = towers[:, 1]
    inFile.z = towers[:, 2]
    inFile.red = colors[:, 0]
    inFile.green = colors[:, 1]
    inFile.blue = colors[:, 2]
    inFile.write('./results/' + fn + '/' + fn + '_towers.las')
    
    #towers2 = np.concatenate(towers2, axis = 0)
    
    #towers2_pcd = o3d.geometry.PointCloud()
    #towers2_pcd.points = o3d.utility.Vector3dVector(towers2)
    #dn_nb = 6
    #dn_std = 3
    #cl, _ = towers2_pcd.remove_statistical_outlier(nb_neighbors=dn_nb, std_ratio=dn_std)  ###6 3
    #towers2 = np.asarray(cl.points)    
    #colors2 = np.zeros_like(towers2) + [[255, 255, 255]]
    ##np.savetxt('./results/' + fn + '/' + fn + '_tower2.xyz', towers2, fmt = '%.6f')
    
    #header = laspy.LasHeader(version="1.4", point_format=7)  # 点格式1对应XYZI，根据你的数据调整格式
    #inFile2 = laspy.LasData(header)
    #inFile2.x = towers2[:, 0]
    #inFile2.y = towers2[:, 1]
    #inFile2.z = towers2[:, 2]
    #inFile2.red = colors2[:, 0]
    #inFile2.green = colors2[:, 1]
    #inFile2.blue = colors2[:, 2]
    #inFile2.write('./results/' + fn + '/' + fn + '_towers2.las')

    #towers_orgs = np.concatenate(towers_orgs, axis = 0)
    #np.savetxt('./results/' + fn + '/' + fn + '_towers_orgs.xyz', towers_orgs, fmt = '%.6f')
    
    #header = laspy.LasHeader(version="1.4", point_format=7)  # 点格式1对应XYZI，根据你的数据调整格式
    #inFile3 = laspy.LasData(header)
    #inFile3.x = towers_orgs[:, 0]
    #inFile3.y = towers_orgs[:, 1]
    #inFile3.z = towers_orgs[:, 2]
    #inFile3.write('./results/' + fn + '/' + fn + '_towers_orgs.las')
    
    return towers_centers

def line_pcds_cal_curvature(line_pcds):
    global progress_value
    print('计算弯曲度',(datetime.datetime.now() - starttime))
    line_curvatures = []
    
    for idx, pcd in enumerate(line_pcds):
        points = np.asarray(pcd.points)
        
        '''
        pcd.estimate_normals()
        #pcd.estimate_normals(search_param=o3d.geometry.KDTreeSearchParamHybrid(radius=15, max_nn=30))
        pcd.orient_normals_consistent_tangent_plane(k=30)
        normals = np.asarray(pcd.normals)
        
        kdtree = KDTree(points)
        _, indices = kdtree.query(points, k = 100)
        
        normal_nerghbors = normals[indices,:]
        normal_dist = np.mean(normal_nerghbors, axis=1)-normals

        curvature = np.mean(np.sqrt(np.sum(normal_dist**2, axis = 1)))
        line_curvatures.append(curvature)
        '''
        
        
        kdtree = KDTree(points)
        _, indices = kdtree.query(points, k=30)
        neighbors = points[indices]
        centered_points = neighbors - points.reshape(-1,1,3)
        cov_matrix = np.matmul(centered_points.transpose(0,2,1), centered_points).astype(np.float32) / (30-1)
        
        U, S, Vt = np.linalg.svd(cov_matrix)    
        normals = U[:,:,0]  
        
        knn = 100
        kdtree = KDTree(points)
        _, indices = kdtree.query(points, k = knn)

        normal_nerghbors = normals[indices,:]
        
        model = normal_nerghbors * normals.reshape(-1,1,3)
        flag_value = np.sum(model, axis = 2)
        flag = np.where(flag_value>0, 1, -1)
        normal_nerghbors = normal_nerghbors * flag.reshape(-1,knn,1)
        
        normal_dist = np.mean(normal_nerghbors, axis=1)-normals   ##mean-dist
        curvature = np.mean(np.sqrt(np.sum(normal_dist**2, axis = 1)))
        line_curvatures.append(curvature)
        
        #normal_dist = np.sqrt(np.sum((normal_nerghbors - normals.reshape(-1,1,3))**2, axis = 2))  #dist-mean
        #curvature = np.mean(normal_dist)
        #line_curvatures.append(curvature)
        
        progress_value = 65 + (68-65) * idx / len(line_pcds)
    line_curvatures = np.array(line_curvatures)
    
    return line_curvatures

def line_pcds_cal_length(line_pcds):
    print('计算长度',(datetime.datetime.now() - starttime))
    line_lengths = []
    
    for idx, pcd in enumerate(line_pcds):
        points = np.asarray(pcd.points)
        
        length = np.sum(np.sqrt(np.sum((points[:-1,:] - points[1:,:])**2, axis = 1)))
        
        line_lengths.append(length)
        
    line_lengths = np.array(line_lengths)
    
    return line_lengths
    
def line_pcds_cal_width(line_pcds):
    print('计算宽度',(datetime.datetime.now() - starttime))
    line_widths = []
    
    for idx, pcd in enumerate(line_pcds):
        points = np.asarray(pcd.points)
        
        start = points[np.argmin(points[:,0]),:]
        end = points[np.argmax(points[:,0]),:]
        
        width = np.sqrt(np.sum((start[:2] - end[:2])**2))
        line_widths.append(width)
        
    line_widths = np.array(line_widths)
    
    return line_widths
    
def line_pcds_cal_height(line_pcds):
    print('计算高度',(datetime.datetime.now() - starttime))
    line_heights = []
    
    for idx, pcd in enumerate(line_pcds):
        points = np.asarray(pcd.points)
        
        height = np.abs(np.max(points[:,2]) - np.min(points[:,2]))
        line_heights.append(height)
        
    line_heights = np.array(line_heights)
    
    return line_heights
    
def detect_danger2(simple_pcd, pcd):
    print('搜索危险物')
    points = np.asarray(pcd.points)
    points_query = np.asarray(simple_pcd.points)
    kdtree = KDTree(points)
    
    dist, indices = kdtree.query(points_query, k = 1)
    
    #danger v1 
    idxs1 = np.where((dist<5) & (dist>1))[0]
    danger_points1 = points_query[idxs1,:]
    colors1 = np.zeros_like(danger_points1) + [255, 0, 0]
    #danger v2 
    idxs2 = np.where((dist<10) & (dist>=5))[0]
    danger_points2 = points_query[idxs2,:]
    colors2 = np.zeros_like(danger_points2) + [255, 255, 0]
    #danger v3 
    idxs3 = np.where((dist<15) & (dist>=10))[0]
    danger_points3 = points_query[idxs3,:]
    colors3 = np.zeros_like(danger_points3) + [0, 255, 0]
    #danger v4 
    idxs4 = np.where((dist<20) & (dist>=15))[0]
    danger_points4 = points_query[idxs4,:]
    colors4 = np.zeros_like(danger_points4) + [0, 0, 255]
    
    danger_points = np.concatenate([danger_points1, danger_points2, danger_points3, danger_points4], axis = 0)
    colors = np.concatenate([colors1, colors2, colors3, colors4], axis = 0)
    
    header = laspy.LasHeader(version="1.4", point_format=7)  # 点格式1对应XYZI，根据你的数据调整格式
    inFile = laspy.LasData(header)
    inFile.x = danger_points[:, 0]
    inFile.y = danger_points[:, 1]
    inFile.z = danger_points[:, 2]
    inFile.red = colors[:,0]
    inFile.green = colors[:,1]
    inFile.blue = colors[:,2]
    
    inFile.write('./results/' + fn + '/' + fn + '_danger.las')
    
def detect_danger(simple_pcd, pcd, towers_centers):
    global fn
    print('搜索危险物')
    import tempfile

    # 确保输出目录存在
    output_dir = os.path.join('results', fn)
    try:
        os.makedirs(output_dir, exist_ok=True)
    except Exception as e:
        print(f'创建输出目录失败: {e}')
        raise

    points = np.asarray(pcd.points)
    points_query = np.asarray(simple_pcd.points)
    kdtree = KDTree(points)

    dist, _ = kdtree.query(points_query, k=1)
    
    # 去除塔
    kdtree_2d = KDTree(towers_centers[:,:2])
    dist_2d, _ = kdtree_2d.query(points_query[:,:2], k=1)

    # Danger v1: 距离在1到5单位之间的点
    idxs1 = np.where((dist < 5) & (dist > 1) & (dist_2d > 15))[0]
    danger_points1 = points_query[idxs1, :]
    colors1 = np.zeros_like(danger_points1) + [255, 0, 0]

    # Danger v2: 距离在5到10单位之间的点
    idxs2 = np.where((dist < 10) & (dist >= 5) & (dist_2d > 15))[0]
    danger_points2 = points_query[idxs2, :]
    colors2 = np.zeros_like(danger_points2) + [255, 255, 0]

    # Danger v3: 距离在10到15单位之间的点
    idxs3 = np.where((dist < 15) & (dist >= 10) & (dist_2d > 15))[0]
    danger_points3 = points_query[idxs3, :]
    colors3 = np.zeros_like(danger_points3) + [0, 255, 0]

    # Danger v4: 距离在15到20单位之间的点
    idxs4 = np.where((dist < 20) & (dist >= 15) & (dist_2d > 15))[0]
    danger_points4 = points_query[idxs4, :]
    colors4 = np.zeros_like(danger_points4) + [0, 0, 255]

    danger_points = np.concatenate([danger_points1, danger_points2, danger_points3, danger_points4], axis=0)
    colors = np.concatenate([colors1, colors2, colors3, colors4], axis=0)

    # 保存危险物点云到LAS文件
    las_output_path = os.path.join('results', fn, f'{fn}_danger.las')
    try:
        header = laspy.LasHeader(version="1.4", point_format=7)
        inFile = laspy.LasData(header)
        inFile.x = danger_points[:, 0]
        inFile.y = danger_points[:, 1]
        inFile.z = danger_points[:, 2]
        inFile.red = colors[:, 0]
        inFile.green = colors[:, 1]
        inFile.blue = colors[:, 2]
        inFile.write(las_output_path)
        print(f'危险物点云已保存至: {las_output_path}')
    except Exception as e:
        print(f'保存LAS文件失败: {e}')
        raise

    # 对每一级危险物点云进行密度聚类并分析
    danger_levels = [
        (danger_points1, '高危 (1-5单位)'),
        (danger_points2, '中危 (5-10单位)'),
        (danger_points3, '低危 (10-15单位)'),
        (danger_points4, '潜在危险 (15-20单位)')
    ]
    level_stats = []

    for points, level_name in danger_levels:
        cluster_count = 0
        volume = 0.0
        avg_height = 0.0

        if len(points) < 10:
            print(f'{level_name} 点云数量不足，无法进行密度聚类')
            level_stats.append((level_name, cluster_count, volume, avg_height))
            continue

        try:
            # 密度聚类
            pcd_temp = o3d.geometry.PointCloud()
            pcd_temp.points = o3d.utility.Vector3dVector(points)
            labels = np.array(pcd_temp.cluster_dbscan(eps=3.5, min_points=10, print_progress=False))
            max_label = labels.max()
            cluster_count = max_label + 1 if max_label >= 0 else 0
            print(f'{level_name} 检测到 {cluster_count} 个危险物')

            # 分析簇的空间范围和高度
            if cluster_count > 0:
                volumes = []
                heights = []
                for cluster_id in range(cluster_count):
                    cluster_indices = np.where(labels == cluster_id)[0]
                    cluster_points = points[cluster_indices]

                    # 空间范围（Bounding Box 体积）
                    if len(cluster_points) > 0:
                        bbox = o3d.geometry.AxisAlignedBoundingBox.create_from_points(
                            o3d.utility.Vector3dVector(cluster_points)
                        )
                        bbox_volume = bbox.volume()
                        volumes.append(bbox_volume)

                        # 平均高度
                        cluster_height = np.mean(cluster_points[:, 2])
                        heights.append(cluster_height)

                # 计算平均值
                volume = np.mean(volumes) if volumes else 0.0
                avg_height = np.mean(heights) if heights else 0.0
                print(f'{level_name} 空间范围: {volume:.2f} 立方米, 平均高度: {avg_height:.2f} 米')

            level_stats.append((level_name, cluster_count, volume, avg_height))

        except Exception as e:
            print(f'{level_name} 密度聚类或分析失败: {e}')
            level_stats.append((level_name, 0, 0.0, 0.0))

    # 生成危险物报告
    report_path = os.path.join('results', fn, f'{fn}_danger_report.docx')
    try:
        doc = Document()
        doc.add_heading('电力线场景危险物分析报告', 0)
        doc.add_paragraph(f'报告生成时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # 统计信息
        doc.add_heading('危险物统计', level=1)
        doc.add_paragraph('以下为检测到的危险物分析结果：')

        table = doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '危险等级'
        hdr_cells[1].text = '危险物数量'
        hdr_cells[2].text = '空间范围（立方米）'
        hdr_cells[3].text = '平均高度（米）'

        for i, (level, cluster_count, volume, avg_height) in enumerate(level_stats, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = level
            row_cells[1].text = str(cluster_count)
            row_cells[2].text = f'{volume:.2f}'
            row_cells[3].text = f'{avg_height:.2f}'
            color = {
                '高危 (1-5单位)': RGBColor(255, 0, 0),
                '中危 (5-10单位)': RGBColor(255, 255, 0),
                '低危 (10-15单位)': RGBColor(0, 255, 0),
                '潜在危险 (15-20单位)': RGBColor(0, 0, 255)
            }[level]
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = color

        # 危险物描述
        doc.add_heading('危险物描述', level=1)
        doc.add_paragraph(
            '1. 高危 (1-5单位)：距离电力线极近，可能为树枝、建筑等，需立即处理以防电力故障。空间范围较大或高度较高可能表示更大规模的障碍物，需优先关注。')
        doc.add_paragraph('2. 中危 (5-10单位)：与电力线有一定距离，可能构成威胁，建议定期检查。')
        doc.add_paragraph('3. 低危 (10-15单位)：较远物体，影响较小，需持续监控。较大的空间范围可能需要额外的巡检计划。')
        doc.add_paragraph(
            '4. 潜在危险 (15-20单位)：距离较远，通常无直接威胁，建议纳入规划。高度较低的物体可能为低矮植被，需关注长期变化。')

        # 处理建议
        doc.add_heading('处理建议', level=1)
        for level, cluster_count, volume, avg_height in level_stats:
            suggestions = []
            if cluster_count > 0:
                suggestions.append(f'- {level}：检测到 {cluster_count} 个危险物。')
                if volume > 100:
                    suggestions.append('  空间范围较大，建议进行全面的现场勘查。')
                if avg_height > 30:
                    suggestions.append('  平均高度较高，可能为高空障碍物，需优先处理以防接触电力线。')
                if not suggestions[1:]:
                    suggestions.append('  定期监控，防止物体进一步靠近电力线。')
            else:
                suggestions.append(f'- {level}：未检测到危险物，建议维持常规监控。')
            for suggestion in suggestions:
                doc.add_paragraph(suggestion)

        # 保存报告，尝试主路径和临时路径
        try:
            doc.save(report_path)
            print(f'危险物报告已保存至: {report_path}')
        except PermissionError as e:
            print(f'保存报告至 {report_path} 失败: {e}')
            temp_dir = tempfile.gettempdir()
            temp_report_path = os.path.join(temp_dir, f'{fn}_danger_report.docx')
            try:
                doc.save(temp_report_path)
                print(f'报告已保存至临时路径: {temp_report_path}')
                report_path = temp_report_path
            except Exception as e:
                print(f'保存报告至临时路径 {temp_report_path} 也失败: {e}')
                # 即使报告保存失败，继续返回 las_output_path
    except Exception as e:
        print(f'生成或保存报告失败: {e}')
    
def generate_catenary(point_A, point_B, point_P, gravity_dir=None, num_points=100):
    """
    根据两个固定点和一个线上任意点生成三维悬链线（P为局部坐标系原点）
    
    参数:
        point_A, point_B: 三维固定点 (numpy数组)
        point_P: 线上任意点 (作为局部坐标系原点)
        gravity_dir: 重力方向向量 (默认 [0,0,-1])
        num_points: 生成点的数量
        
    返回:
        numpy数组: [num_points, 3] 悬链线点坐标
    """
    # 默认重力方向为负z轴
    if gravity_dir is None:
        gravity_dir = np.array([0, 0, -1])
    gravity_dir = gravity_dir / np.linalg.norm(gravity_dir)

    # 计算平面法向量（AP × BP）
    AP = point_A - point_P
    BP = point_B - point_P
    #normal = np.cross(AP, BP)
    normal = np.cross(point_A-point_B, gravity_dir)  #默认垂直
    if np.linalg.norm(normal) < 1e-10:
        print("三点共线，无法确定平面, 默认垂直")
        #return np.array([point_A, point_B, point_P])
        normal = np.cross(point_A-point_B, gravity_dir)
    normal = normal / np.linalg.norm(normal)

    # 计算y轴（重力反方向的平面内投影）
    gravity_up = -gravity_dir
    proj = np.dot(gravity_up, normal) * normal
    y_axis = gravity_up - proj
    if np.linalg.norm(y_axis) < 1e-10:
        print("重力方向与平面法向平行")
        return np.array([point_A, point_B, point_P])
    y_axis = y_axis / np.linalg.norm(y_axis)

    # 计算x轴（平面内与y轴正交）
    x_axis = np.cross(y_axis, normal)
    x_axis = x_axis / np.linalg.norm(x_axis)

    # 转换A、B到局部坐标系（P为原点）
    A_local_x = np.dot(AP, x_axis)
    A_local_y = np.dot(AP, y_axis)
    B_local_x = np.dot(BP, x_axis)
    B_local_y = np.dot(BP, y_axis)

    # 求解悬链线参数a和x0（约束：P(0,0)在曲线上）
    def equations(params):
        a, x0 = params
        a = max(a, 1e-5)  # 避免除零
        # 约束条件：A、B在曲线上，且P(0,0)满足y=0
        eq1 = a * np.cosh((A_local_x - x0)/a) - a - A_local_y
        eq2 = a * np.cosh((B_local_x - x0)/a) - a - B_local_y
        eq3 = a * np.cosh(-x0/a) - a  # P点y坐标强制为0
        return np.array([eq1, eq2, eq3])

    # 初始猜测：a为两点高度差的均值，x0为A/B的x坐标均值
    a_initial = np.mean([abs(A_local_y), abs(B_local_y)])
    x0_initial =  np.mean([A_local_x, B_local_x])
    
    #params = fsolve(equations, [a_initial, x0_initial])
    #a_val, x0_val = params[0], params[1]
    
    result = root(equations, [a_initial, x0_initial], method='lm')  # Levenberg-Marquardt算法
    a_val, x0_val = result.x[0], result.x[1]

    # 生成局部坐标系中的点
    x_min = min(A_local_x, B_local_x, 0) - 0.5
    x_max = max(A_local_x, B_local_x, 0) + 0.5
    x_range = np.linspace(x_min, x_max, num_points)
    y_local = a_val * np.cosh((x_range - x0_val)/a_val) - a_val

    # 转换回全局坐标系
    points_global = []
    for i in range(num_points):
        point = point_P + x_range[i] * x_axis + y_local[i] * y_axis
        points_global.append(point)
    
    return np.array(points_global)
    
def line_pcds_subdiv(line_pcds):
    global progress_value
    print('细分，计算半径',(datetime.datetime.now() - starttime))
    line_vectors_all = []
    line_radii = []
    for idx, pcd in tqdm(enumerate(line_pcds)):
    
        points = np.asarray(pcd.points).copy()
        colors = np.asarray(pcd.colors)
        
        start_idx = np.argmin(points[:,0])
        start = points[start_idx,:]
        end_idx = np.argmax(points[:,0])
        end = points[end_idx,:]

        
        kdtree = KDTree(points)
        
        _, c2_idx =kdtree.query(np.array([start, end]), k = 30)      
        c2_neighbors = points[c2_idx,]
        c2_neighbors_mean = np.mean(c2_neighbors, axis = 1)
        
        start = c2_neighbors_mean[0,:]
        end = c2_neighbors_mean[1,:]
        
        ###########
        idxs_seg = np.where((points[:,0] > ((end[0] - start[0]) * 0.2 + start[0])) & (points[:,0] < ((end[0] - start[0]) * 0.8 + start[0])))[0]
        points_seg = points[idxs_seg,:]
        
        _, idxs =kdtree.query(points_seg, k = 30)
        neighbors = points[idxs,]
        neighbors_mean = np.mean(neighbors, axis = 1)
        
        points_straight = points_seg.copy()
        
        points_straight[:,2] = points_straight[:,2] - neighbors_mean[:,2]
        
        #np.savetxt('./results/' + fn + './line_org_' + str(idx) +'.xyz', points, fmt = '%.6f')
        #np.savetxt('./results/' + fn + './line_curvature_' + str(idx) +'.xyz', points_seg, fmt = '%.6f')
        #np.savetxt('./results/' + fn + './line_straight_' + str(idx) +'.xyz', points_straight, fmt = '%.6f')
        
        start_idx_seg = np.argmin(points_seg[:,0])
        start_seg = points_seg[start_idx_seg,:]
        end_idx_seg = np.argmax(points_seg[:,0])
        end_seg = points_seg[end_idx_seg,:]
        
        _, c2_idx_seg =kdtree.query(np.array([start_seg, end_seg]), k = 30)
        c2_neighbors_seg = points[c2_idx_seg,]
        c2_neighbors_mean_seg = np.mean(c2_neighbors_seg, axis = 1)
        
        start_seg = c2_neighbors_mean_seg[0,:]
        end_seg = c2_neighbors_mean_seg[1,:]
        
        direction = start_seg-end_seg
        direction[2]= 0
        direction = direction / np.linalg.norm(direction)
        
        org = (start_seg+end_seg)/2
        org[2] = 0
        
        point_vec = points_straight - org.reshape(-1,3)
        dot_product = np.dot(point_vec, direction)
        projected_points = points_straight - np.outer(dot_product, direction)
        
        #projected_points = projected_points - np.mean(projected_points, axis =0)
        #np.savetxt('./results/' + fn + './line_straight_projection_' + str(idx) + '.xyz', projected_points + np.array([[1,1,0]]) * idx, fmt = '%.6f')
        
        projected_pcd = o3d.geometry.PointCloud()
        projected_pcd.points = o3d.utility.Vector3dVector(projected_points)
        
        dn_nb = 6
        dn_std = 1
        filter_pcd, _ = projected_pcd.remove_statistical_outlier(nb_neighbors=dn_nb, std_ratio=dn_std)  ###6 3
        
        flter_points = np.asarray(filter_pcd.points)
        
        box = np.max(flter_points, axis = 0) - np.min(flter_points, axis = 0)
        width = np.sqrt(box[0]**2 + box[1]**2)
        height = box[2]
        
        labels = None
        
        #if(width>0.9 and height> 0.65):
        #    cls_num = 6
        #    
        #    kmeans = KMeans(n_clusters= cls_num)  
        #    labels = kmeans.fit_predict(flter_points)

        if(height < 0.2 and width < 0.5):
            cls_num = 1
            labels = np.zeros([flter_points.shape[0]])

        else:
            flter_points_z = flter_points.copy()
            if(width<0.5):
                flter_points_z[:,0] = 0
                flter_points_z[:,1] = 0   
        
            sil_scores = []
            labels_all = []
            for k in [2,3,4,5,6]:
                kmeans = KMeans(n_clusters=k, random_state=42)
                labels_temp = kmeans.fit_predict(flter_points_z)
                sil_scores.append(silhouette_score(flter_points_z, labels_temp))
                labels_all.append(labels_temp)
            
            optimal_k = np.argmax(sil_scores)
            cls_num = optimal_k + 2
            #print(idx, ':', cls_num, sil_scores)
            labels = labels_all[optimal_k]
            
            #if(height > 0.65 and cls_num==2):  #高2转4
            #    cls_num = 4
            #    labels = labels_all[4-2]
            #elif((height < 0.65 and cls_num!=2) or (cls_num==2 and sil_scores[optimal_k] < 0.7)):  #低多/低置信度 转1
            #    cls_num = 1
            #    labels[:] = 0
                
        #filter_pcd_cls = np.concatenate([flter_points + np.array([[1,1,0]]) * idx,labels.reshape(-1,1)], axis =1)
        #np.savetxt('./results/' + fn + '/line_straight_projection_kmeans_cluster_'+ str(idx) + '.xyz', filter_pcd_cls, fmt = '%.6f')
        
        center = np.mean(flter_points, axis = 0)
        vectors = []
        radii = []
        
        if(cls_num==1):
            box_temp = np.max(flter_points, axis = 0) - np.min(flter_points, axis = 0)
            radius = (np.sqrt(box_temp[0]**2 + box_temp[1]**2) + box_temp[2])/2/2
            radii.append(radius)
            vectors.append(np.array([0,0,0]))
        else:
            for i in range(cls_num):
                idxs_cluster = np.where(labels==i)[0]
                points_cluster = flter_points[idxs_cluster,:]
                box_temp = np.max(points_cluster, axis = 0) - np.min(points_cluster, axis = 0)
                radius = (np.sqrt(box_temp[0]**2 + box_temp[1]**2) + box_temp[2])/2/2
                radii.append(radius)
                center_cluster = np.mean(points_cluster, axis = 0) 
                vector_cluster = center_cluster - center
                vectors.append(vector_cluster)
        vectors = np.array(vectors)
        radius_mean = np.mean(np.array(radii))
            
        line_vectors_all.append(vectors)
        line_radii.append(radius_mean)
        
        progress_value = 68 + (98-68) * idx / len(line_pcds)
    line_radii = np.array(line_radii)
    return line_vectors_all, line_radii

def quadratic_func(X, a, b, c, d, e, f):
    x, y = X
    return a + b*x + c*y + d*x**2 + e*x*y + f*y**2
    
def line_curve_fit(points, start, end, density):
    initial_guess = [1, 1, 1, 0.1, 0.1, 0.1]  # 初始参数猜测
    
    center = np.mean(points, axis = 0)
    points = points - center.reshape(1,3)
    popt, pcov = curve_fit(quadratic_func, (points[:,0], points[:,1]), points[:,2], p0=initial_guess)
    
    alpha = np.arange(0,density,1)/density
    init_points = start.reshape(-1,3) + (end - start).reshape(-1,3) * alpha.reshape(-1,1) - center
    init_points[:,2] = quadratic_func((init_points[:,0], init_points[:,1]), *popt)
    
    reconstruction_points = init_points + center
    
    return reconstruction_points
    
def line_pcds_reconstruction(line_pcds, density = 1000):
    print('重建',(datetime.datetime.now() - starttime))
    line_restruction_pcds = []
    for idx, pcd in tqdm(enumerate(line_pcds)):
        #dn_nb = 6
        #dn_std = 3
        #pcd, _ = pcd.remove_statistical_outlier(nb_neighbors=dn_nb, std_ratio=dn_std)  ###6 3
    
        points = np.asarray(pcd.points)
        colors = np.asarray(pcd.colors)
        
        start_idx = np.argmin(points[:,0])
        start = points[start_idx,:]
        end_idx = np.argmax(points[:,0])
        end = points[end_idx,:]
        
        mid_idx = np.argmin(points[:,2])
        mid = points[mid_idx,:]
        
        dist1 = np.linalg.norm(start-mid)
        dist2 = np.linalg.norm(end-mid)
        
        if((dist1 < 3) or (dist2 < 3)):
            #print('无最低点，使用中间点')
            kdtree = KDTree(points)
            _, mid_idx = kdtree.query(((start + end)/2).reshape(1,3), k=1)
            mid = points[mid_idx[0],:]
        
        kdtree = KDTree(points)
        
        _, c3_idx =kdtree.query(np.array([start, end, mid]), k = 10) # 10 
        c3_neighbors = points[c3_idx,]
        c3_neighbors_mean = np.mean(c3_neighbors, axis = 1)
        start = c3_neighbors_mean[0,:]
        end = c3_neighbors_mean[1,:]
        mid = c3_neighbors_mean[2,:]
        
        reconstruction_colors = np.zeros([density, 3]) + colors[0,:].reshape(-1,3)
        
        reconstruction_points = line_curve_fit(points, start, end, density)
        
        #reconstruction_points = generate_catenary(start, end, mid, gravity_dir=[0,0,-1], num_points=density)
        
        #pro = np.array([[1,1,0]])
        #kdtree2 = KDTree(points * pro)
        
        if(len(reconstruction_points)==3):
            #print('拟合失败，再重建')
            
            alpha = np.arange(0,density,1)/density
            init_points = start.reshape(-1,3) * alpha.reshape(-1,1) + (end - start).reshape(-1,3)
            dist, idxs = kdtree.query(init_points, k=10)
            neighbors = points[idxs,]
            neighbors_mean = np.mean(neighbors, axis = 1)
            reconstruction_points_temp = init_points.copy()
            reconstruction_points_temp = neighbors_mean
            
            kdtree3 = KDTree(reconstruction_points_temp)
            dist, idxs2 = kdtree3.query(init_points, k=30)
            reconstruction_points = np.mean(reconstruction_points_temp[idxs2], axis = 1)
            
        else:
            dist, _ = kdtree.query(reconstruction_points, k=1)
            dist_max = np.max(dist)
            if(dist_max>2):
                #print('拟合误差过大，再重建')
                
                init_points = reconstruction_points
                dist, idxs = kdtree.query(init_points, k=10)
                neighbors = points[idxs,]
                neighbors_mean = np.mean(neighbors, axis = 1)
                reconstruction_points_temp = init_points.copy()
                reconstruction_points_temp = neighbors_mean
                
                kdtree3 = KDTree(reconstruction_points_temp)
                dist, idxs2 = kdtree3.query(init_points, k=30)
                reconstruction_points = np.mean(reconstruction_points_temp[idxs2], axis = 1)
                
        
        line_restruction_pcd = o3d.geometry.PointCloud()
        line_restruction_pcd.points = o3d.utility.Vector3dVector(reconstruction_points)
        line_restruction_pcd.colors = o3d.utility.Vector3dVector(reconstruction_colors)

        line_restruction_pcds.append(line_restruction_pcd)
        
    return line_restruction_pcds
    
def output_pcds(line_pcds):    
    global progress_value
    line_restruction_pcds = line_pcds_reconstruction(line_pcds, 3000)
    
    line_idxs = np.arange(0, len(line_pcds), 1) * 10
    line_lengths = line_pcds_cal_length(line_restruction_pcds)
    line_widths = line_pcds_cal_width(line_restruction_pcds)    
    line_heights = line_pcds_cal_height(line_pcds)
    line_curvatures = line_pcds_cal_curvature(line_pcds)
    line_vectors_all, line_radii =  line_pcds_subdiv(line_pcds)
    
    progress_value = 98
    print('存储',(datetime.datetime.now() - starttime))
    ###output excel 
    output_excel = np.concatenate([line_idxs.reshape(-1,1), line_lengths.reshape(-1,1), line_widths.reshape(-1,1), line_heights.reshape(-1,1), line_curvatures.reshape(-1,1), line_radii.reshape(-1,1)], axis = 1)
    np.savetxt('./results/' + fn + '/' + fn + '_results_table.csv', output_excel, delimiter = ',', fmt = '%.6f')
    
    
    ###output origin point
    pc_points_all = []
    pc_colors_all = []
    pc_idxs_all = []    
    pc_lengths_all = []
    pc_widths_all = []
    pc_heights_all = []
    pc_curvatures_all = []
    pc_radii_all = []
    
    for idx, pcd in enumerate(line_pcds):
        line_length = line_lengths[idx]
        line_width = line_widths[idx]
        line_height = line_heights[idx]
        line_curvature = line_curvatures[idx]
        line_radius = line_radii[idx]
        
        
        points = np.asarray(pcd.points)
        
        pc_points_all.extend(points)
        pc_colors_all.extend(np.asarray(pcd.colors))
        pc_idxs_all.extend(np.ones([points.shape[0]])* idx * 10)
        pc_lengths_all.extend(np.ones([points.shape[0]])* line_length)
        pc_widths_all.extend(np.ones([points.shape[0]])* line_width)
        pc_heights_all.extend(np.ones([points.shape[0]])* line_height)
        pc_curvatures_all.extend(np.ones([points.shape[0]])* line_curvature)
        pc_radii_all.extend(np.ones([points.shape[0]])* line_radius)

        
    pc_points_all = np.array(pc_points_all)
    pc_colors_all = np.array(pc_colors_all) * 255
    pc_colors_all = pc_colors_all.astype(np.int32)
    pc_idxs_all = np.array(pc_idxs_all).astype(np.int32)
    pc_lengths_all = np.array(pc_lengths_all)
    pc_widths_all = np.array(pc_widths_all)
    pc_heights_all = np.array(pc_heights_all)
    pc_curvatures_all = np.array(pc_curvatures_all)
    pc_radii_all = np.array(pc_radii_all)

    header = laspy.LasHeader(version="1.4", point_format=7)  # 点格式1对应XYZI，根据你的数据调整格式
    custom_fields = [
    {"name": "idx", "type": np.int32, "description": "编号"},
    {"name": "length", "type": np.float32, "description": "长度"},
    {"name": "width", "type": np.float32, "description": "宽度"},
    {"name": "height", "type": np.float32, "description": "高度"},    
    {"name": "curvature", "type": np.float32, "description": "弯曲度"},
    {"name": "radius", "type": np.float32, "description": "半径"}]
    for field in custom_fields:
        header.add_extra_dim(laspy.ExtraBytesParams(**field)) 
    
    inFile = laspy.LasData(header)
    inFile.x = pc_points_all[:, 0]
    inFile.y = pc_points_all[:, 1]
    inFile.z = pc_points_all[:, 2]
    inFile.red = pc_colors_all[:, 0]
    inFile.green = pc_colors_all[:, 1]
    inFile.blue = pc_colors_all[:, 2]
    
    user_attrs = {"idx": pc_idxs_all, "length": pc_lengths_all, "width": pc_widths_all, "height": pc_heights_all, "curvature": pc_curvatures_all, "radius": pc_radii_all}
    
    for key, value in user_attrs.items():
        if hasattr(inFile, key):
            setattr(inFile, key, value)
        elif hasattr(inFile.header, key):
            setattr(inFile.header, key, value)
    
    inFile.write('./results/' + fn + '/' + fn + '_results_orgin.las')
    
    ###output reconstruction point
    pc_points_all = []
    pc_colors_all = []
    pc_idxs_all = []    
    pc_lengths_all = []
    pc_widths_all = []
    pc_heights_all = []
    pc_curvatures_all = []
    pc_radii_all = []
    
    for idx, line_restruction_pcd in enumerate(line_restruction_pcds):
        line_length = line_lengths[idx]
        line_width = line_widths[idx]
        line_height = line_heights[idx]
        line_curvature = line_curvatures[idx]
        line_radius = line_radii[idx]
        line_vectors = line_vectors_all[idx]
        
        line_restruction_points = np.asarray(line_restruction_pcd.points)
        line_restruction_colors = np.asarray(line_restruction_pcd.colors)
        
        pc_idxs = np.ones([line_restruction_points.shape[0]])* idx * 10
        pc_lengths = np.ones([line_restruction_points.shape[0]])* line_length
        pc_widths = np.ones([line_restruction_points.shape[0]])* line_width
        pc_heights = np.ones([line_restruction_points.shape[0]])* line_height
        pc_curvatures = np.ones([line_restruction_points.shape[0]])* line_curvature
        pc_radii = np.ones([line_restruction_points.shape[0]])* line_radius
        
        for sub_idx in range(len(line_vectors)):
            pc_points_all.extend(line_restruction_points + line_vectors[sub_idx])
            pc_colors_all.extend(line_restruction_colors)
            pc_idxs_all.extend(pc_idxs + sub_idx)
            pc_lengths_all.extend(pc_lengths)
            pc_widths_all.extend(pc_widths)
            pc_heights_all.extend(pc_heights)
            pc_curvatures_all.extend(pc_curvatures)
            pc_radii_all.extend(pc_radii)

    pc_points_all = np.array(pc_points_all)
    pc_colors_all = np.array(pc_colors_all) * 255
    pc_colors_all = pc_colors_all.astype(np.int32)
    pc_idxs_all = np.array(pc_idxs_all).astype(np.int32)
    pc_lengths_all = np.array(pc_lengths_all)
    pc_widths_all = np.array(pc_widths_all)
    pc_heights_all = np.array(pc_heights_all)    
    pc_curvatures_all = np.array(pc_curvatures_all)
    pc_radii_all = np.array(pc_radii_all)     
    
    header = laspy.LasHeader(version="1.4", point_format=7)  # 点格式1对应XYZI，根据你的数据调整格式
    custom_fields = [
    {"name": "idx", "type": np.int32, "description": "编号"},
    {"name": "length", "type": np.float32, "description": "长度"},
    {"name": "width", "type": np.float32, "description": "宽度"},
    {"name": "height", "type": np.float32, "description": "高度"},    
    {"name": "curvature", "type": np.float32, "description": "弯曲度"},
    {"name": "radius", "type": np.float32, "description": "半径"}]
    for field in custom_fields:
        header.add_extra_dim(laspy.ExtraBytesParams(**field)) 
    
    inFile = laspy.LasData(header)
    inFile.x = pc_points_all[:, 0]
    inFile.y = pc_points_all[:, 1]
    inFile.z = pc_points_all[:, 2]
    inFile.red = pc_colors_all[:, 0]
    inFile.green = pc_colors_all[:, 1]
    inFile.blue = pc_colors_all[:, 2]
    
    user_attrs = {"idx": pc_idxs_all, "length": pc_lengths_all, "width": pc_widths_all, "height": pc_heights_all, "curvature": pc_curvatures_all, "radius": pc_radii_all}
    
    for key, value in user_attrs.items():
        if hasattr(inFile, key):
            setattr(inFile, key, value)
        elif hasattr(inFile.header, key):
            setattr(inFile.header, key, value)
    
    inFile.write('./results/' + fn + '/' + fn + '_results_reconstruction.las')
    
def process_path(path):
    try:
        print(path)
        if not os.path.exists(path):
            print('file ' + path + ' does not exist!')
            return -1
        
        global progress_value
        progress_value = 0
        
        global starttime
        starttime = datetime.datetime.now()
        
        global fn 
        fn = path.split('/')[-1].split('\\')[-1][:-4]
        
        print(path, fn)
        
        if not os.path.exists('./results/'):
            os.mkdir('./results/')
        
        if not os.path.exists('./results/' + fn + '/'):
            os.mkdir('./results/' + fn + '/')
        
        print('Strat, time',(datetime.datetime.now() - starttime))

        progress_value = 1
        print('step 1/8: load point cloud, time:',(datetime.datetime.now() - starttime))
        
        point_pcd = read_file(path)
        
        progress_value = 5
        print('step 2/8: filter point cloud, time:',(datetime.datetime.now() - starttime))
        
        simple_pcd = simple(point_pcd)#采样
        
        progress_value = 10
        print('step 3/8: extract candidate points, time:',(datetime.datetime.now() - starttime))
        
        red_pcd = rough_extraction(simple_pcd)#提取线
        
        progress_value = 48
        print('step 4/8: extract candidate lines, time:',(datetime.datetime.now() - starttime))
        
        pcds = DBSCAN(red_pcd)#聚类
        
        progress_value = 50
        print('step 5/8: extract fine lines, time:',(datetime.datetime.now() - starttime))
        
        pcds = maxpoints(pcds)
        
        progress_value = 55
        print('step 6/8: select lines, time:',(datetime.datetime.now() - starttime))
        
        pcd, line_pcds, towers_centers = is_line(pcds, simple_pcd)#提取线
        
        progress_value = 56
        print('step 7/8: detect danger object, time:',(datetime.datetime.now() - starttime))
        
        detect_danger(simple_pcd, pcd, towers_centers)

        progress_value = 65
        print('step 8/8: calculate features and reconstruct, time:',(datetime.datetime.now() - starttime))
        
        output_pcds(line_pcds)
        
        progress_value = 100
        print('Finsh, time:',(datetime.datetime.now() - starttime))
        
        return '123'

    except Exception as e:
        raise Exception(f"Error processing point cloud: {str(e)}")


if __name__ == '__main__':
    path = sys.argv[1]
    
    process_path(path)
    
    
    
    
    
    
    
    